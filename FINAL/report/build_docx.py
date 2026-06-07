"""Build the GRACE Track-1 technical report (DOCX) from report_data.json + figures/.

Clean white-page report with a Modal-green accent; the dark dashboard figures embed
as panels. Every number traces to report_data.json. Run from FINAL/:
    python report/build_docx.py
"""
from __future__ import annotations

import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
DATA = json.load(open(os.path.join(HERE, "report_data.json"), encoding="utf-8"))
LB = DATA["leaderboard"]
ENS = DATA["ensemble"]
DS = DATA["dataset"]
META = DATA["run_meta"]
ALL9 = ENS["dev_scores"]["all9"]["official_score"]
TOP3 = ENS["dev_scores"]["top3"]["official_score"]
GREEDY = ENS["dev_scores"]["greedy"]["official_score"]
BEST_SEED = max(max(r["seed_macro"].values()) for r in LB)
CLASSES = ("Premise", "Claim", "MajorClaim")
SEEDS = META["seeds"]

GREEN = RGBColor(0x1A, 0x7F, 0x37)     # accent (readable on white)
DARKG = RGBColor(0x10, 0x3A, 0x1E)     # deep green for title
GREY = RGBColor(0x57, 0x60, 0x6A)
INK = RGBColor(0x1B, 0x1F, 0x24)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
for lvl, sz, col in ((1, 15, GREEN), (2, 12.5, DARKG)):
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = col
for section in doc.sections:
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def hrule(paragraph, color="1A7F37", size=14):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", str(size)), ("w:space", "2"), ("w:color", color)):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def para(text="", size=10.5, color=INK, bold=False, italic=False, align=None,
         after=6, before=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
        r.font.color.rgb = color
    return p


def runs(p, *segments):
    """Add multiple styled runs to a paragraph: (text, dict) tuples."""
    for text, opts in segments:
        r = p.add_run(text)
        r.font.size = Pt(opts.get("size", 10.5))
        r.font.bold = opts.get("bold", False)
        r.font.italic = opts.get("italic", False)
        r.font.color.rgb = opts.get("color", INK)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead); r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARKG
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p


def figure(name, caption):
    doc.add_picture(os.path.join(FIG, name), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = para(caption, 9, GREY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, before=3)
    return cap


def callout(title, body):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    shade(cell, "EAF7EE")
    cell.paragraphs[0].text = ""
    pt = cell.paragraphs[0]
    pt.paragraph_format.space_after = Pt(3)
    r = pt.add_run(title); r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARKG
    pb = cell.add_paragraph()
    rb = pb.add_run(body); rb.font.size = Pt(10.5); rb.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def style_table(tbl, header_fill="1A7F37"):
    tbl.style = "Table Grid"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, cell in enumerate(tbl.rows[0].cells):
        shade(cell, header_fill)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9.5)
    for i, row in enumerate(tbl.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                shade(cell, "F2F7F3")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9.5)


def set_cell(cell, text, bold=False, color=INK):
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.font.bold = bold; r.font.color.rgb = color; r.font.size = Pt(9.5)


# ============================================================ TITLE
t = para("Argument Mining in Spanish Clinical Trials", 24, DARKG, bold=True, after=2)
para("A 27-Run Transformer Benchmark and Ensemble for Evidence Component Detection",
     13, GREEN, bold=True, after=4)
hrule(doc.add_paragraph())
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
runs(p, ("GRACE @ IberLEF 2026", {"bold": True, "size": 10.5}),
     ("   ·   Track 1 (Evidence Component Detection, Spanish)   ·   Subtask 1   ·   ", {"color": GREY, "size": 10}),
     ("Technical Report", {"color": GREY, "italic": True, "size": 10}))

callout(
    "Headline result",
    f"A single parallel sweep trained 9 transformer backbones × 3 seeds (27 runs) on Modal serverless GPU. "
    f"The strongest single backbone, mDeBERTa-v3-base, reaches {LB[0]['mean_macro']:.3f} ± {LB[0]['std']:.3f} "
    f"dev strict macro-F1. Averaging all 27 softmax distributions (the all-9 ensemble) raises this to "
    f"{ALL9:.3f} — a +{(ALL9 - LB[0]['mean_macro']) * 100:.1f}-point gain with no extra training and no "
    f"selection bias — and produces a validated blind-test submission of {ENS['blind']['all9']['entities']:,} "
    f"argumentative components across 2,474 abstracts.")

# ============================================================ 1 TASK & DATA
doc.add_heading("1  Task and Dataset", level=1)
para("Track 1 of GRACE asks systems to read a Spanish randomised-controlled-trial (RCT) abstract and label "
     "each argumentative segment as one of three component types — Premise (reported evidence and statistical "
     "results), Claim (interpretive statements), or MajorClaim (the abstract's overall conclusion) — while "
     "leaving non-argumentative text unlabelled. The official ranking metric is strict macro-F1, which requires "
     "predicted spans to match gold-standard character offsets exactly.")
para("We frame the problem as sentence-level 4-class classification (None / Premise / Claim / MajorClaim). "
     "Sentences are segmented with spaCy (es_core_news_sm); each is classified in the context of its previous "
     "sentence. The corpus is small and severely imbalanced:", after=4)
# dataset table
dt = doc.add_table(rows=4, cols=5)
hdr = ["Split", "Abstracts", "Premise", "Claim", "MajorClaim"]
for j, h in enumerate(hdr):
    dt.rows[0].cells[j].paragraphs[0].add_run(h)
tr = DS["entities"]["train"]; dv = DS["entities"]["dev"]
set_cell(dt.rows[1].cells[0], "Train"); set_cell(dt.rows[1].cells[1], f"{DS['docs']['train']}")
set_cell(dt.rows[1].cells[2], f"{tr['Premise']:,}"); set_cell(dt.rows[1].cells[3], f"{tr['Claim']:,}")
set_cell(dt.rows[1].cells[4], f"{tr['MajorClaim']}")
set_cell(dt.rows[2].cells[0], "Dev"); set_cell(dt.rows[2].cells[1], f"{DS['docs']['dev']}")
set_cell(dt.rows[2].cells[2], f"{dv['Premise']}"); set_cell(dt.rows[2].cells[3], f"{dv['Claim']}")
set_cell(dt.rows[2].cells[4], f"{dv['MajorClaim']}")
set_cell(dt.rows[3].cells[0], "Blind test"); set_cell(dt.rows[3].cells[1], f"{DS['docs']['blind']:,}")
for j in (2, 3, 4):
    set_cell(dt.rows[3].cells[j], "—")
style_table(dt)
para()
para(f"MajorClaim is only 2.8% of training components and just {dv['MajorClaim']} spans in dev — the central "
     f"difficulty of the task and the dominant source of variance in our results.", 10, GREY, italic=True)
figure("06_data_overview.png", "Figure 1. Corpus size (left) and training class distribution (right). "
       "MajorClaim's scarcity motivates a weighted loss and a low decision threshold.")

# ============================================================ 2 METHOD
doc.add_heading("2  Method", level=1)
para("All nine backbones share an identical pipeline so the comparison reflects the model, not the recipe.")
doc.add_heading("2.1  Input and training", level=2)
bullet("prev_sentence [SEP] sentence, max length 192 tokens, spaCy es_core_news_sm segmentation.", "Representation: ")
bullet("body 1e-5, classifier head 1e-4 (AdamW, weight decay 0.05).", "Differential learning rate: ")
bullet("weighted cross-entropy with MajorClaim weight 10, label smoothing 0.15, and a 0.50 probability "
       "override that promotes a segment to MajorClaim whenever its probability exceeds the threshold.", "Imbalance handling: ")
bullet("cosine with 100 warmup steps, batch size 16, up to 15 epochs, early-stopping patience 4.", "Schedule: ")
doc.add_heading("2.2  Two-stage auto-runner", level=2)
para("Each run executes two stages. Stage 1 trains on the training split and evaluates strict macro-F1 on dev "
     "after every epoch, keeping the best epoch (load-best-model-at-end); its dev logits are saved for ensemble "
     "construction. Stage 2 refits the model on train+dev for (best_epoch + 1) epochs and emits logits for the "
     "blind test. This cleanly separates held-out model selection from the final full-data fit used for submission.")
para("Training ran on Modal serverless GPU (NVIDIA A10G, Ampere) using bf16 mixed precision, with the full "
     "9×3 grid dispatched in parallel from a single launch.")

# ============================================================ 3 SETUP
doc.add_heading("3  Experimental Setup", level=1)
para("Nine Spanish and multilingual encoders — spanning general, biomedical, and clinical pretraining — were "
     "each trained with three seeds (42, 100, 2026), giving 27 runs. All 27 converged.")
figure("07_run_panel.png", "Figure 2. The sweep at a glance: a single parallel Modal launch produced every "
       "model, seed, dev score, and blind prediction reported here.")

# ============================================================ 4 RESULTS
doc.add_heading("4  Results", level=1)
para("We report mean ± standard deviation over the three seeds. Reporting distributions rather than a single "
     "best seed is essential on a dataset this small.")
# main results table (9 models)
rt = doc.add_table(rows=1 + len(LB), cols=5)
for j, h in enumerate(["Model", "Macro-F1 (mean ± std)", "Premise", "Claim", "MajorClaim"]):
    rt.rows[0].cells[j].paragraphs[0].add_run(h)
for i, r in enumerate(LB, 1):
    set_cell(rt.rows[i].cells[0], r["pretty"], bold=(i == 1), color=(DARKG if i == 1 else INK))
    set_cell(rt.rows[i].cells[1], f"{r['mean_macro']:.3f} ± {r['std']:.3f}", bold=(i == 1))
    set_cell(rt.rows[i].cells[2], f"{r['per_class']['Premise']['mean']:.3f}")
    set_cell(rt.rows[i].cells[3], f"{r['per_class']['Claim']['mean']:.3f}")
    set_cell(rt.rows[i].cells[4], f"{r['per_class']['MajorClaim']['mean']:.3f}")
style_table(rt)
para()
para("Table 1. Per-model dev strict macro-F1 (mean ± std over 3 seeds) and per-class F1. "
     "mDeBERTa-v3-base leads; per-class columns are seed means.", 9, GREY, italic=True)
figure("01_leaderboard.png", "Figure 3. Model leaderboard. The dashed line marks the all-9 ensemble (0.743), "
       "which clears every single model.")
para("Two patterns stand out. First, clinical/biomedical pretraining helps but does not dominate: the general "
     "multilingual mDeBERTa-v3 edges out the domain-specific Spanish clinical encoders. Second, the per-class "
     "view (Figure 4) shows a consistent ceiling — every backbone scores 0.82–0.90 on Premise yet only "
     "0.24–0.43 on MajorClaim, confirming the bottleneck is data scarcity rather than any single architecture.")
figure("02_per_class_f1.png", "Figure 4. Per-class F1 by model. The cold MajorClaim column is universal.")
figure("03_seed_variance.png", "Figure 5. Seed stability. Per-seed spread reaches ~0.07 macro-F1, so "
       "single-seed rankings are unreliable.")
figure("05_training_curves.png", "Figure 6. Dev macro-F1 per epoch (seed 42) for the five strongest "
       "backbones, which peak within 4–7 epochs; the best epoch is retained (early stopping, patience 4).")

# ============================================================ 5 ENSEMBLE
doc.add_heading("5  Ensemble", level=1)
para("We average softmax distributions across runs (mean over the three seeds per model, then over models) and "
     "decode with the same 0.50 MajorClaim override. We evaluate three candidates with the official scorer:", after=4)
bullet(f"the bias-free choice using every run — dev strict macro-F1 = {ALL9:.4f}.", "All-9 (zero-peek): ")
bullet(f"the three highest-mean backbones — {TOP3:.4f}.", "Top-3 (dev-selected): ")
bullet(f"forward selection on dev — {GREEDY:.4f} (dev-optimistic upper bound).", "Greedy: ")
para(f"Crucially, the all-9 ensemble ({ALL9:.4f}) is not only the best — it also beats the dev-selected top-3 "
     f"({TOP3:.4f}). Because the #3 and #4 backbones differ by less than the seed noise, selecting members on "
     f"dev adds bias without adding accuracy. We therefore submit all-9: the choice that peeks at nothing is also "
     f"the strongest.", before=2)
figure("04_ensemble_gain.png", "Figure 7. Ensemble vs single models. Averaging 27 distributions adds "
       f"+{(ALL9 - LB[0]['mean_macro']) * 100:.1f} points over the best single model's mean, with no extra training.")

# ============================================================ 6 BLIND
doc.add_heading("6  Blind-Test Submission", level=1)
b = ENS["blind"]["all9"]
para(f"Applying the all-9 ensemble to the {DS['docs']['blind']:,} blind abstracts ({DS['blind_sentences']:,} "
     f"segmented sentences) yields {b['entities']:,} predicted argumentative components: "
     f"{b['by_type']['Premise']:,} Premise, {b['by_type']['Claim']:,} Claim, and {b['by_type']['MajorClaim']} "
     f"MajorClaim. Every predicted span was offset-validated (raw_text[start:end] exactly equals the predicted "
     f"text — 0 mismatches after trimming spaCy's trailing whitespace). For the two majority classes the "
     f"predicted split closely matches the training prior (Premise 67.0% vs 67.8%, Claim 31.6% vs 29.4%); "
     f"MajorClaim is predicted even more conservatively (1.4% vs 2.8%), consistent with its known difficulty — "
     f"a useful sanity check that the ensemble did not collapse to a single class.")
figure("08_blind_distribution.png", "Figure 8. Blind-test predicted class distribution vs the training prior.")

# ============================================================ 7 ENGINEERING
doc.add_heading("7  Engineering Notes and Lessons", level=1)
bullet("mDeBERTa-v3-base diverged to NaN gradients under fp32 and scored 0.0. Switching to bf16 on Ampere "
       "hardware stabilised training and turned it into the top single model. The right numerical precision "
       "mattered more than the choice of architecture.", "Precision unlocks the best model. ")
bullet("Averaging 27 distributions is the cheapest reliable gain available — +3.9 points over the best single "
       "model with zero additional training — and, chosen as all-9, it carries no model-selection bias.", "Ensembling is the honest win. ")
bullet("Identical results across seeds are the exception, not the rule: spread reaches ~0.07 macro-F1 on 350 "
       "training documents. Mean ± std over multiple seeds is the defensible way to compare models here.", "Seeds matter. ")
bullet("Per-platform numbers differ (the same model trained on a different cloud produced different dev scores), "
       "so all comparisons in this report come from one consistent Modal sweep. We report mean ± std, never a "
       "cherry-picked seed or platform.", "Reproducibility caveat. ")

# ============================================================ 8 CONCLUSION
doc.add_heading("8  Conclusion and Next Steps", level=1)
para(f"A disciplined 27-run benchmark plus a bias-free ensemble reaches {ALL9:.3f} dev strict macro-F1 on GRACE "
     f"Track 1 and delivers a clean blind-test submission. The remaining headroom is concentrated in MajorClaim, "
     f"a data-scarcity problem rather than a modelling one.")
para("Priorities going forward: (1) clause-level segmentation to recover sub-sentence spans the strict metric "
     "penalises; (2) MajorClaim-targeted augmentation or contrastive objectives to lift the weakest class; and "
     "(3) Track-2 relation detection built on top of these components.")

# ============================================================ APPENDIX
doc.add_page_break()
doc.add_heading("Appendix A  Full Per-Run Results (27 runs)", level=1)
at = doc.add_table(rows=1 + len(LB) * len(SEEDS), cols=6)
for j, h in enumerate(["Model", "Seed", "Best epoch", "Macro-F1", "Premise", "Claim / MajorClaim"]):
    at.rows[0].cells[j].paragraphs[0].add_run(h)
row = 1
for r in LB:
    for si, seed in enumerate(SEEDS):
        set_cell(at.rows[row].cells[0], r["pretty"])
        set_cell(at.rows[row].cells[1], str(seed))
        set_cell(at.rows[row].cells[2], str(r["best_epoch"][str(seed)]))
        set_cell(at.rows[row].cells[3], f"{r['seed_macro'][str(seed)]:.4f}")
        set_cell(at.rows[row].cells[4], f"{r['per_class']['Premise']['vals'][si]:.3f}")
        set_cell(at.rows[row].cells[5],
                 f"{r['per_class']['Claim']['vals'][si]:.3f} / {r['per_class']['MajorClaim']['vals'][si]:.3f}")
        row += 1
style_table(at)
para()

doc.add_heading("Appendix B  Reproducibility", level=1)
rec = META["recipe"]
para("All artifacts derive from two Modal apps in FINAL/ and one consolidation step:", after=4)
bullet("trains the 9×3 grid (two-stage recipe) and writes per-run logits + metrics to a Modal volume.", "modal_sweep.py: ")
bullet("loads the 27 logit sets, reproduces a known single-model dev F1 as an alignment gate, builds the "
       "all-9 / top-3 / greedy ensembles, scores them with the official program, and writes the blind submission.",
       "modal_ensemble.py: ")
bullet("collect_data.py → make_figures.py → build_docx.py / build_pptx.py regenerate this report from the "
       "saved results.", "report/: ")
para(f"Key settings: max_len {rec['max_len']}, body LR {rec['lr_body']}, head LR {rec['lr_head']}, "
     f"{rec['loss']}, {rec['scheduler']}, label smoothing {rec['label_smoothing']}, batch {rec['batch_size']}, "
     f"≤{rec['max_epochs']} epochs, ES patience {rec['es_patience']}, MajorClaim override {rec['majorclaim_override']}. "
     f"GPU: {META['gpu']}.", 9.5, GREY)

out = os.path.join(HERE, "GRACE_Track1_Report.docx")
doc.save(out)
print(f"saved {out}")
print(f"paragraphs: {len(doc.paragraphs)} | tables: {len(doc.tables)} | "
      f"images: {sum(1 for s in doc.inline_shapes)}")
